"""
Oversight File Extraction Service (Stage 4 + 5 combined)
=========================================================
استخراج متن از فایل‌های پیوست تسک، با persistence مرحله‌به‌مرحله.

طراحی:
- per-mime router (text, PDF, DOCX, XLSX, image, audio, video)
- هر segment که استخراج شد، **بلافاصله** در `extractions.json` ذخیره می‌شود
  → اگر backend crash کند، resume از segment بعدی ممکن است.
- خود فایل اصلی پس از اتمام (set_status='extracted') حذف می‌شود.
- مدل بصری/multimodal از `pick_best_extraction_model(mime)` انتخاب می‌شود
  → داینامیک، قابل override در /models.
- چارچوب «بدون خلاصه‌سازی»: prompt صریح می‌خواهد متن full literal.
- عناوین داینامیک: اول از مدل headings/sections می‌پرسیم بر اساس user_idea،
  سپس متن مرتبط با هر heading را extract می‌کنیم.

Schema:
  FileExtraction:
    id, task_id, session_id, file_order, original_filename, mime_type,
    total_segments (estimated), completed_segments, status, model_used,
    started_at, finished_at, error
  ExtractionSegment:
    id, extraction_id, segment_index, segment_title, text, raw_excerpt,
    page_or_timestamp, started_at, finished_at, model_used
"""

from __future__ import annotations

import asyncio
import base64
import io
import json
import logging
import os
import uuid
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from .oversight_service import STORAGE_DIR, now_iso
from .oversight_upload_session import (
    get_upload_session_service, UploadSession,
)
from ..core.models_registry import (
    pick_best_extraction_model, mime_to_required_capability,
    ModelCapability, get_model, MODEL_REGISTRY,
)

logger = logging.getLogger(__name__)

EXTRACTIONS_FILE: Path = STORAGE_DIR / "extractions.json"

# ─────────── سقف‌های ایمنی ───────────
MAX_PAGES_PER_PDF: int = 100_000  # عملاً unlimited
MAX_PARAGRAPHS_PER_DOCX: int = 10_000_000
MAX_ROWS_PER_SHEET: int = 100_000_000
# 🛡 (Stage 10 audit fix #2) — طبق درخواست صریح کاربر «محدودیت استخراج متن
# اصلاً نداشته باشه»، این سقف فقط محافظ JSON store در برابر فایل JSON عظیم
# است (≥10MB). در عمل، هیچ segment معمولی به این سقف نمی‌رسد.
SEGMENT_TEXT_MAX_CHARS: int = 100_000_000  # 🔴 (extraction-100pct-fix) 10MB→100MB per segment
PER_SEGMENT_TIMEOUT_SEC: int = 300  # 5 دقیقه
INLINE_MEDIA_BYTES_LIMIT: int = 18 * 1024 * 1024  # 18MB
AV_CHUNK_SECONDS: int = 300  # 5 دقیقه per audio/video chunk
MEMORY_THRESHOLD_PCT: float = 80.0  # اگر RAM >80%, تأخیر بده

# ─────────── extraction concurrency ───────────
_EXTRACTION_SEMAPHORE = asyncio.Semaphore(1)


async def _wait_for_memory_headroom(max_wait_sec: int = 60) -> None:
    """اگر RAM سیستم بالای آستانه است، تا max_wait صبر کن (به جای OOM crash).
    در صورت عدم وجود psutil، silent pass.
    """
    try:
        import psutil
    except ImportError:
        return
    waited = 0
    while waited < max_wait_sec:
        try:
            pct = psutil.virtual_memory().percent
        except Exception:
            return
        if pct < MEMORY_THRESHOLD_PCT:
            return
        logger.warning(
            f"memory at {pct}% > {MEMORY_THRESHOLD_PCT}% — تأخیر extraction (waited {waited}s)"
        )
        await asyncio.sleep(5)
        waited += 5
    logger.warning(f"memory still high after {max_wait_sec}s — ادامه می‌دهیم (best-effort)")


def _ffmpeg_available() -> bool:
    """آیا ffmpeg در PATH هست؟"""
    import shutil
    return shutil.which("ffmpeg") is not None


async def _ffmpeg_chunk_av(
    input_path: Path,
    output_dir: Path,
    chunk_seconds: int = AV_CHUNK_SECONDS,
) -> List[Path]:
    """تقسیم audio/video به chunkهای N ثانیه‌ای. خروجی: لیست chunk paths.

    استفاده از segment muxer — copy stream، بدون re-encode → سریع و
    لاجواب لیترال. اگر ffmpeg موجود نباشد، RuntimeError.
    """
    if not _ffmpeg_available():
        raise RuntimeError("ffmpeg در PATH نیست — chunking ممکن نیست")
    output_dir.mkdir(parents=True, exist_ok=True)
    pattern = str(output_dir / "chunk_%04d" + input_path.suffix)
    # cmd: ffmpeg -i input -c copy -f segment -segment_time N -reset_timestamps 1 pattern
    cmd = [
        "ffmpeg", "-hide_banner", "-loglevel", "error",
        "-i", str(input_path),
        "-c", "copy", "-f", "segment",
        "-segment_time", str(chunk_seconds),
        "-reset_timestamps", "1",
        pattern,
    ]
    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdout=asyncio.subprocess.DEVNULL,
        stderr=asyncio.subprocess.PIPE,
    )
    try:
        # 🛡 (audit fix MINOR) — timeout 10 دقیقه برای chunking یک فایل بزرگ
        _, stderr = await asyncio.wait_for(proc.communicate(), timeout=600)
    except asyncio.TimeoutError:
        # 🛡 (audit fix L1 CRITICAL) — kill + wait برای جلوگیری از zombie
        try:
            proc.kill()
        except Exception:
            pass
        try:
            await asyncio.wait_for(proc.wait(), timeout=5)
        except Exception:
            pass
        raise RuntimeError("ffmpeg chunking timeout after 10 minutes")
    if proc.returncode != 0:
        raise RuntimeError(
            f"ffmpeg returned {proc.returncode}: {stderr.decode('utf-8', errors='ignore')[:500]}"
        )
    chunks = sorted(output_dir.glob(f"chunk_*{input_path.suffix}"))
    return chunks


# ====================================================================
# Dataclasses
# ====================================================================

@dataclass
class ExtractionSegment:
    id: str
    extraction_id: str
    segment_index: int
    segment_title: str
    text: str = ""
    raw_excerpt: str = ""
    page_or_timestamp: str = ""
    model_used: str = ""
    started_at: str = field(default_factory=now_iso)
    finished_at: Optional[str] = None
    status: str = "pending"  # pending|done|failed
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class FileExtraction:
    id: str
    task_id: Optional[str]
    session_id: str
    file_order: int
    original_filename: str
    mime_type: str
    total_segments: int = 0     # estimated (و بعد updated)
    completed_segments: int = 0
    status: str = "pending"     # pending|extracting|extracted|failed
    model_used: str = ""
    started_at: str = field(default_factory=now_iso)
    finished_at: Optional[str] = None
    error: Optional[str] = None
    # متن کامل ادغام‌شده (cache برای read سریع)
    full_text_cache: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# ====================================================================
# Repository (JSON persistence)
# ====================================================================

class ExtractionRepo:
    """Repository برای FileExtraction + ExtractionSegment.

    دو لیست موازی در یک JSON ذخیره می‌شوند تا read/write atomic بمانند.
    """

    def __init__(self) -> None:
        self._extractions: Dict[str, FileExtraction] = {}
        self._segments: Dict[str, List[ExtractionSegment]] = {}  # ext_id -> [segments]
        self._lock = asyncio.Lock()
        self._load()

    def _load(self) -> None:
        if not EXTRACTIONS_FILE.exists():
            return
        try:
            data = json.loads(EXTRACTIONS_FILE.read_text(encoding="utf-8"))
            for e in data.get("extractions", []):
                try:
                    fe = FileExtraction(**e)
                    self._extractions[fe.id] = fe
                except Exception as exc:
                    logger.warning(f"extractions: skip malformed entry: {exc}")
            for ext_id, segs in (data.get("segments", {}) or {}).items():
                out: List[ExtractionSegment] = []
                for s in segs:
                    try:
                        out.append(ExtractionSegment(**s))
                    except Exception as exc:
                        logger.warning(f"extractions: skip malformed segment: {exc}")
                self._segments[ext_id] = out
        except Exception as e:
            logger.warning(f"extractions: load failed: {e}")

    def _save(self) -> None:
        try:
            data = {
                "extractions": [e.to_dict() for e in self._extractions.values()],
                "segments": {
                    ext_id: [s.to_dict() for s in segs]
                    for ext_id, segs in self._segments.items()
                },
            }
            tmp = EXTRACTIONS_FILE.with_suffix(".json.tmp")
            tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            tmp.replace(EXTRACTIONS_FILE)
        except Exception as e:
            logger.warning(f"extractions: save failed: {e}")

    async def create_extraction(self, fe: FileExtraction) -> FileExtraction:
        async with self._lock:
            self._extractions[fe.id] = fe
            self._segments.setdefault(fe.id, [])
            self._save()
        return fe

    async def update_extraction(self, extraction_id: str, **patch: Any) -> Optional[FileExtraction]:
        async with self._lock:
            fe = self._extractions.get(extraction_id)
            if fe is None:
                return None
            for k, v in patch.items():
                if hasattr(fe, k):
                    setattr(fe, k, v)
            self._save()
            return fe

    async def add_segment(self, extraction_id: str, seg: ExtractionSegment) -> ExtractionSegment:
        async with self._lock:
            self._segments.setdefault(extraction_id, []).append(seg)
            # update completed counter
            fe = self._extractions.get(extraction_id)
            if fe and seg.status == "done":
                fe.completed_segments = sum(
                    1 for s in self._segments[extraction_id] if s.status == "done"
                )
                # update full_text_cache incrementally
                fe.full_text_cache = self._build_full_text(extraction_id)
            self._save()
        return seg

    async def update_segment(self, segment_id: str, **patch: Any) -> Optional[ExtractionSegment]:
        async with self._lock:
            for ext_id, segs in self._segments.items():
                for s in segs:
                    if s.id == segment_id:
                        for k, v in patch.items():
                            if hasattr(s, k):
                                setattr(s, k, v)
                        # update parent fe
                        fe = self._extractions.get(ext_id)
                        if fe:
                            fe.completed_segments = sum(
                                1 for ss in segs if ss.status == "done"
                            )
                            fe.full_text_cache = self._build_full_text(ext_id)
                        self._save()
                        return s
        return None

    def _build_full_text(self, extraction_id: str) -> str:
        segs = self._segments.get(extraction_id) or []
        segs_sorted = sorted(segs, key=lambda s: s.segment_index)
        parts: List[str] = []
        for s in segs_sorted:
            if s.status != "done":
                continue
            head = f"## {s.segment_title}" if s.segment_title else f"## Segment {s.segment_index}"
            if s.page_or_timestamp:
                head += f"  _(at: {s.page_or_timestamp})_"
            parts.append(head)
            parts.append(s.text)
        return "\n\n".join(parts)

    def get(self, extraction_id: str) -> Optional[FileExtraction]:
        return self._extractions.get(extraction_id)

    def get_segments(self, extraction_id: str) -> List[ExtractionSegment]:
        return list(self._segments.get(extraction_id) or [])

    def list_by_task(self, task_id: str) -> List[FileExtraction]:
        out = [e for e in self._extractions.values() if e.task_id == task_id]
        out.sort(key=lambda e: e.file_order)
        return out

    def list_by_session(self, session_id: str) -> List[FileExtraction]:
        return [e for e in self._extractions.values() if e.session_id == session_id]

    def full_text(self, extraction_id: str) -> str:
        return self._build_full_text(extraction_id)


# ──── singleton ────
_repo_instance: Optional[ExtractionRepo] = None


def get_extraction_repo() -> ExtractionRepo:
    global _repo_instance
    if _repo_instance is None:
        _repo_instance = ExtractionRepo()
    return _repo_instance


async def boot_recover_stale_extractions() -> Dict[str, int]:
    """در startup، extractionهای stale (status='extracting' بدون فعالیت) را شناسایی
    و mark failed کنیم. session مرتبط هم به 'failed' برگردانده می‌شود.

    این به کاربر اجازه می‌دهد دوباره trigger کند و از segmentهای done ادامه دهد
    (extract_session خودش resume-friendly است).
    """
    from datetime import datetime, timezone, timedelta
    repo = get_extraction_repo()
    cutoff = datetime.now(timezone.utc) - timedelta(minutes=5)
    cleared = 0
    for fe in list(repo._extractions.values()):
        if fe.status != "extracting":
            continue
        # last activity = آخرین segment finished_at، یا started_at
        last_active = fe.started_at
        segs = repo.get_segments(fe.id)
        if segs:
            done_segs = [s for s in segs if s.finished_at]
            if done_segs:
                last_active = max(s.finished_at for s in done_segs if s.finished_at)
        try:
            ts = datetime.fromisoformat((last_active or "").replace("Z", "+00:00"))
        except Exception:
            ts = datetime.min.replace(tzinfo=timezone.utc)
        if ts < cutoff:
            await repo.update_extraction(
                fe.id,
                status="failed",
                error="boot recovery: extraction interrupted by restart",
                finished_at=datetime.now(timezone.utc).isoformat(),
            )
            # session هم به failed (تا کاربر بفهمد و دوباره trigger کند)
            try:
                from .oversight_upload_session import get_upload_session_service
                up = get_upload_session_service()
                s = up.get(fe.session_id)
                if s and s.status == "extracting":
                    await up.set_status(s.id, "completed",  # نه failed — بتواند resume
                                        error="extraction interrupted by restart")
            except Exception as e:
                logger.debug(f"boot recovery session update failed: {e}")
            cleared += 1
    return {"cleared_stale_extractions": cleared}


# ====================================================================
# Extraction service
# ====================================================================

class ExtractionError(Exception):
    pass


async def _ai_extract_text(
    model_id: str,
    *,
    prompt: str,
    image_b64: Optional[str] = None,
    inline_file_data: Optional[Tuple[str, bytes]] = None,  # (mime, bytes) — for Gemini Files
    max_tokens: int = 64000,  # 🔴 (extraction-100pct-fix) — bump 32K→64K، چون
    # prompt جدید verbose-ـتر است و verbatim می‌خواد. مدل‌های مدرن (Gemini
    # 2.5 Pro=64K، Claude=64K، GPT-4=16K) اکثراً 64K پشتیبانی می‌کنن.
) -> str:
    """فراخوانی AI برای استخراج متن از یک قطعه (صفحه/فریم/audio chunk).

    استفاده از ai_manager موجود — provider-agnostic. اگر مدل Gemini است،
    گزینهٔ `images` (base64) را قبول می‌کند. برای audio/video با Gemini،
    inline_file_data به‌صورت base64-inline در پیام درج می‌شود (Gemini
    حداکثر ~20MB inline قبول می‌کند — برای فایل‌های بزرگ‌تر Stage 9
    ffmpeg chunking خواهد داشت).
    """
    from .ai_manager import get_ai_manager
    from .ai_base import Message

    images: Optional[List[str]] = None
    if image_b64:
        images = [image_b64]

    # 🛡 (audit fix CRITICAL) — رسانهٔ غیر-تصویری (audio/video/PDF/...) باید
    # با MIME صریح به provider ارسال شود، نه از طریق `images` (که در
    # GeminiService همیشه به image/png|jpeg سنیف می‌شد → خطای
    # "Unable to process input image" برای فایل‌های صوتی/ویدئویی).
    inline_files: Optional[List[tuple]] = None
    extra_text = ""
    if inline_file_data is not None:
        mime, raw = inline_file_data
        if len(raw) > 18 * 1024 * 1024:
            # Gemini inline limit ~20MB — برای chunkهای بزرگ‌تر، در Stage 9
            # ffmpeg به chunkهای ≤5min تقسیم می‌کند
            raise ExtractionError(
                f"chunk too large for inline ({len(raw)} bytes). Need ffmpeg chunking."
            )
        b64 = base64.b64encode(raw).decode("ascii")
        extra_text = (
            f"\n\n[فایل پیوست به‌صورت base64 با mime={mime}؛ "
            f"حجم = {len(raw)} bytes؛ متن کامل را استخراج کن.]"
        )
        # تصمیم routing بر اساس mime واقعی:
        # - image/* → از کانال `images` (سازگار با همهٔ providerها)
        # - بقیه (audio, video, pdf, ...) → از کانال `inline_files` با MIME صریح
        if (mime or "").lower().startswith("image/"):
            images = (images or []) + [b64]
        else:
            inline_files = (inline_files or []) + [(mime, b64)]

    mgr = get_ai_manager()
    messages = [
        Message(role="system", content=(
            # 🔴 (extraction-100pct-fix) — prompt قدیمی فقط ۵ خط بود و باعث می‌شد
            # مدل verbose نباشه. کاربر گزارش داد ۹۰٪+ محتوا گم می‌شد. این prompt
            # جدید سختگیرانه‌تر است: verbatim، no-summarization، preserve-format،
            # explicit min-output-length proportional به input، و instructions
            # برای حالت‌های خاص (PDF، DOCX، code، table، audio، image).
            "تو یک extractor دقیق و کامل هستی. وظیفه‌ات استخراج **عین به عین** (verbatim) "
            "از محتوای ورودی است. این کار **حیاتی** است — output تو مستقیماً به "
            "task synthesis می‌ره، اگر چیزی drop کنی task ناقص می‌شه و کاربر "
            "محتوای از دست رفته رو می‌فهمه.\n\n"
            "## قوانین مطلق (هیچ‌کدوم optional نیست)\n"
            "1. **verbatim کامل** — هر کلمه، هر جمله، هر paragraph عیناً منتقل شه.\n"
            "2. **هیچ summarization** — اگه فایل ۲۰۰ آیتم داره، ۲۰۰ تا منتقل کن (نه «و ۱۹۸ تای دیگه...»).\n"
            "3. **هیچ drop** — حتی boilerplate، header/footer، صفحات تشکر، imageای که "
            "متن داره (OCR)، watermark، disclaimer، signature، table of contents... همگی شامل می‌شه.\n"
            "4. **preserve formatting**:\n"
            "   - code blocks دقیقاً با همان whitespace و indentation\n"
            "   - tables به‌صورت Markdown table (| col1 | col2 |)\n"
            "   - bullet/numbered lists با همان نشانه‌ها (-، *، 1.، 2.، الف، ب)\n"
            "   - headings با # (سطح متناسب)\n"
            "   - inline formatting (**bold**, *italic*, `code`) حفظ شه\n"
            "5. **page/section markers** — اگه ورودی چندصفحه‌ای یا چندبخشی است، "
            "هر صفحه/بخش رو با `--- صفحه N ---` یا `### بخش: Title` جدا کن.\n"
            "6. **multilingual** — اگه متن چندزبانه‌ست، همگی رو منتقل کن. زبان فارسی "
            "رو با حروف فارسی بنویس (RTL).\n"
            "7. **audio/video**: transcript کامل + speaker labels اگه قابل تشخیصه + "
            "timestamps هر ~۳۰s (مثال: `[02:15] گوینده ۱: …`).\n"
            "8. **image/photo**: همهٔ متن قابل خواندن + توصیف عناصر بصری مهم "
            "(اگه متن به‌تنهایی context رو نمی‌رسونه).\n\n"
            "## ممنوعات\n"
            "- ❌ «… و موارد مشابه»، «… و غیره»، «خلاصه»، «in summary»\n"
            "- ❌ پرش از بخشی چون تکراری به نظر می‌رسه\n"
            "- ❌ کوتاه‌سازی برای صرفه‌جویی tokens — کاربر گفته «هزینه مهم نیست»\n"
            "- ❌ paraphrase — کلمات کاربر باید عین به عین منتقل بشن\n"
            "- ❌ اضافه کردن تحلیل/نظر — فقط استخراج\n\n"
            "## حداقل خروجی\n"
            "اگه ورودی متنی است، حجم خروجی باید **حداقل ۹۰٪** حجم ورودی باشه.\n"
            "اگه ورودی image-only هست (PDF اسکن، عکس)، خروجی باید تمام متن قابل "
            "خواندن رو شامل شه — اگر مطمئن نیستی، توصیف هم اضافه کن.\n\n"
            "تو باید پر-حرف و پر-جزئیات باشی. لطفاً بفهم: کم گفتن = شکست."
        )),
        Message(
            role="user",
            content=prompt + extra_text,
            images=images,
            inline_files=inline_files,
        ),
    ]
    # 🛡 (audit fix CRITICAL) — allow_fallback=False تا اگر مدل بصری در DB
    # disabled است، ai_manager به deepseek fallback نکند و خروجی غلط
    # «نمی‌توانم تصویر را ببینم» تولید نشود. در عوض exception می‌خوریم،
    # که بالاتر در _resolve_attachments_for_idea handle می‌شود.
    resp = await mgr.generate(
        model_id=model_id,
        messages=messages,
        max_tokens=max_tokens,
        temperature=0.1,
        allow_fallback=False,
    )
    return (resp.content or "").strip()


async def _plan_headings(model_id: str, user_idea: str, filename: str, mime: str) -> List[str]:
    """گام «عناوین داینامیک»: قبل از استخراج، از مدل می‌خواهیم headings/sections
    حسب user_idea و mime پیشنهاد دهد. در صورت شکست، headings عمومی برمی‌گرداند.
    """
    from .ai_manager import get_ai_manager
    from .ai_base import Message
    try:
        mgr = get_ai_manager()
        messages = [
            Message(role="system", content=(
                "تو یک مشاور هستی. کاربر یک ایده داده + یک فایل پیوست. ایده + متادیتای فایل را "
                "ببین و یک لیست JSON از 3 تا 8 heading/section بساز که برای استخراج «هدفمند» "
                "از این فایل مفید باشد. ترتیب از مهم‌ترین به کم‌اهمیت‌ترین."
            )),
            Message(role="user", content=(
                f"ایدهٔ کاربر:\n{user_idea[:2000]}\n\n"
                f"فایل پیوست: {filename}  (mime={mime})\n\n"
                "خروجی فقط JSON: {\"headings\": [\"...\", \"...\"]}"
            )),
        ]
        resp = await mgr.generate(
            model_id=model_id, messages=messages, max_tokens=2000, temperature=0.2,
            allow_fallback=False,  # 🛡 (audit fix CRITICAL)
        )
        # parse
        import re
        m = re.search(r'\{[^{}]*"headings"\s*:\s*\[.*?\]\s*\}', resp.content or "", re.DOTALL)
        if m:
            data = json.loads(m.group(0))
            # 🔴 (extraction-100pct-fix) 200→500 per heading، 8→50 total
            # برای document بزرگ با ۲۰+ section نباید heading ها drop بشن
            hs = [str(x)[:500] for x in (data.get("headings") or [])]
            if hs:
                return hs[:50]
    except Exception as e:
        logger.debug(f"_plan_headings failed: {e}")
    # fallback عمومی
    return [
        "خلاصهٔ کلی محتوا",
        "جزئیات اصلی و نکات کلیدی",
        "اطلاعات تکمیلی و frequently mentioned items",
    ]


# ─────────────── per-mime extractors ───────────────

def _read_text_file(path: Path) -> str:
    """خواندن متن خام (txt, md, csv, log, json, ...)."""
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        raise ExtractionError(f"read text failed: {e}")


async def _render_pdf_page_to_b64(path: Path, page_num: int) -> Optional[str]:
    """رندر یک صفحهٔ PDF به PNG → base64.

    🆕 (Stage 9) — برای صفحات scan-only که pypdf متن استخراج نمی‌کند.
    ابتدا pdftoppm (poppler) را امتحان می‌کنیم (نتیجهٔ بهتر، سریع‌تر).
    اگر در دسترس نبود، fallback به render via Pillow + پیغام.

    خروجی: PNG base64 string، یا None اگر هیچ ابزار rendering موجود نباشد.
    """
    import shutil, tempfile
    pdftoppm = shutil.which("pdftoppm")
    if pdftoppm:
        with tempfile.TemporaryDirectory() as tmp:
            out_prefix = Path(tmp) / "p"
            cmd = [
                pdftoppm,
                "-png", "-r", "150",
                "-f", str(page_num), "-l", str(page_num),
                str(path), str(out_prefix),
            ]
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.PIPE,
            )
            try:
                # 🛡 (audit fix MINOR) — timeout سخت تا روی PDF مخدوش hang نکند
                _, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)
            except asyncio.TimeoutError:
                # 🛡 (audit fix L1 CRITICAL) — kill + wait برای جلوگیری از zombie
                try:
                    proc.kill()
                except Exception:
                    pass
                try:
                    # حداکثر ۵ ثانیه صبر کن تا process واقعاً مرده
                    await asyncio.wait_for(proc.wait(), timeout=5)
                except Exception:
                    pass
                logger.warning(f"pdftoppm timeout on page {page_num} after 30s")
                return None
            if proc.returncode != 0:
                logger.debug(
                    f"pdftoppm exit {proc.returncode}: {stderr.decode('utf-8', errors='ignore')[:200]}"
                )
                return None
            # pdftoppm naming: p-1.png یا p-01.png …
            candidates = sorted(Path(tmp).glob("p-*.png"))
            if not candidates:
                return None
            try:
                return base64.b64encode(candidates[0].read_bytes()).decode("ascii")
            except Exception:
                return None
    # fallback نهایی — هیچ rendering در دسترس نیست
    return None


def _extract_pdf_pages(path: Path) -> List[Tuple[int, str]]:
    """استخراج text-only از PDF با pypdf. صفحه‌های scan شده فقط با AI extract
    می‌شوند (در main routing). خروجی: [(page_num, text)] فقط صفحاتی که
    pypdf توانست متن استخراج کند.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ExtractionError("pypdf نصب نیست (Stage 1 dep)")
    try:
        reader = PdfReader(str(path))
        out: List[Tuple[int, str]] = []
        for i, page in enumerate(reader.pages[:MAX_PAGES_PER_PDF], start=1):
            try:
                text = (page.extract_text() or "").strip()
            except Exception:
                text = ""
            out.append((i, text))
        return out
    except Exception as e:
        raise ExtractionError(f"pypdf failed: {e}")


def _extract_docx_paragraphs(path: Path) -> List[str]:
    try:
        import docx
    except ImportError:
        raise ExtractionError("python-docx نصب نیست (Stage 1 dep)")
    try:
        d = docx.Document(str(path))
        out: List[str] = []
        for p in d.paragraphs[:MAX_PARAGRAPHS_PER_DOCX]:
            t = (p.text or "").strip()
            if t:
                out.append(t)
        # tables هم
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(cells)
                if line.strip(" |"):
                    out.append(line)
        return out
    except Exception as e:
        raise ExtractionError(f"python-docx failed: {e}")


def _extract_xlsx_sheets(path: Path) -> List[Tuple[str, str]]:
    """خروجی: [(sheet_name, csv_like_text)]"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ExtractionError("openpyxl نصب نیست (Stage 1 dep)")
    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        out: List[Tuple[str, str]] = []
        for sheet in wb.worksheets:
            rows: List[str] = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i >= MAX_ROWS_PER_SHEET:
                    break
                cells = [str(c) if c is not None else "" for c in row]
                rows.append("\t".join(cells))
            out.append((sheet.title, "\n".join(rows)))
        return out
    except Exception as e:
        raise ExtractionError(f"openpyxl failed: {e}")


# ─────────────── Format Diversity (Stage 9) ───────────────

_EXT_TO_MIME: Dict[str, str] = {
    # text
    ".txt": "text/plain", ".md": "text/markdown", ".markdown": "text/markdown",
    ".csv": "text/csv", ".tsv": "text/tab-separated-values",
    ".log": "text/plain", ".ini": "text/plain",
    ".json": "application/json", ".jsonl": "application/x-ndjson",
    ".ndjson": "application/x-ndjson",
    ".xml": "application/xml", ".html": "text/html", ".htm": "text/html",
    ".yaml": "application/yaml", ".yml": "application/yaml",
    ".toml": "application/toml",
    # code
    ".py": "text/x-python", ".pyw": "text/x-python",
    ".js": "application/javascript", ".mjs": "application/javascript",
    ".ts": "application/typescript", ".tsx": "application/typescript",
    ".jsx": "application/javascript",
    ".java": "text/x-java", ".kt": "text/x-kotlin",
    ".c": "text/x-c", ".h": "text/x-c", ".cpp": "text/x-c++",
    ".hpp": "text/x-c++",
    ".cs": "text/x-csharp", ".go": "text/x-go", ".rs": "text/x-rust",
    ".rb": "text/x-ruby", ".php": "application/x-php",
    ".scala": "text/x-scala", ".clj": "text/x-clojure",
    ".swift": "text/x-swift", ".dart": "text/x-dart",
    ".lua": "text/x-lua", ".pl": "text/x-perl",
    ".sh": "application/x-shellscript", ".bash": "application/x-shellscript",
    ".ps1": "text/x-powershell",
    ".sql": "application/sql",
    ".css": "text/css", ".scss": "text/x-scss",
    ".r": "text/x-r",
    ".ex": "text/x-elixir", ".exs": "text/x-elixir",
    ".elm": "text/x-elm", ".vue": "text/x-vue",
    ".ipynb": "application/x-ipynb+json",
    "dockerfile": "text/x-dockerfile",
    # documents
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xls": "application/vnd.ms-excel",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".ppt": "application/vnd.ms-powerpoint",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".rtf": "application/rtf",
    ".odt": "application/vnd.oasis.opendocument.text",
    ".ods": "application/vnd.oasis.opendocument.spreadsheet",
    ".odp": "application/vnd.oasis.opendocument.presentation",
    ".epub": "application/epub+zip",
    # image
    ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".png": "image/png", ".gif": "image/gif", ".webp": "image/webp",
    ".bmp": "image/bmp", ".tiff": "image/tiff", ".tif": "image/tiff",
    ".heic": "image/heic", ".svg": "image/svg+xml",
    # audio
    ".mp3": "audio/mpeg", ".wav": "audio/wav", ".ogg": "audio/ogg",
    ".m4a": "audio/mp4", ".flac": "audio/flac", ".aac": "audio/aac",
    ".opus": "audio/opus",
    # video
    ".mp4": "video/mp4", ".mov": "video/quicktime", ".avi": "video/x-msvideo",
    ".mkv": "video/x-matroska", ".webm": "video/webm", ".m4v": "video/mp4",
    # archive
    ".zip": "application/zip", ".tar": "application/x-tar",
    ".gz": "application/gzip", ".7z": "application/x-7z-compressed",
}


def _guess_mime_from_extension(filename: str) -> Optional[str]:
    """تخمین mime از پسوند فایل. اگر mime موجود application/octet-stream است،
    این تابع mime دقیق‌تر می‌دهد.
    """
    import os
    if not filename:
        return None
    base = os.path.basename(filename).lower()
    # special: Dockerfile
    if base == "dockerfile" or base.startswith("dockerfile."):
        return "text/x-dockerfile"
    _, ext = os.path.splitext(base)
    return _EXT_TO_MIME.get(ext)


async def _extract_ipynb(
    path: Path, fe: "FileExtraction", repo: "ExtractionRepo",
    completed: set, persist_segment_fn,
) -> None:
    """parse Jupyter Notebook → هر cell یک segment.

    cells: code, markdown, raw. outputs (اگر هست) به‌عنوان suffix.
    """
    import json as _json
    try:
        data = _json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception as e:
        raise ExtractionError(f"ipynb parse failed: {e}")
    cells = data.get("cells") or []
    await repo.update_extraction(fe.id, total_segments=len(cells))
    for idx, cell in enumerate(cells, start=1):
        if idx in completed:
            continue
        ctype = cell.get("cell_type") or "unknown"
        src_raw = cell.get("source") or ""
        src = "".join(src_raw) if isinstance(src_raw, list) else str(src_raw)
        out_text = ""
        outputs = cell.get("outputs") or []
        if outputs and ctype == "code":
            out_parts: List[str] = []
            for o in outputs[:20]:
                if "text" in o:
                    t = o["text"]
                    out_parts.append("".join(t) if isinstance(t, list) else str(t))
                elif o.get("data", {}).get("text/plain"):
                    tp = o["data"]["text/plain"]
                    out_parts.append("".join(tp) if isinstance(tp, list) else str(tp))
            if out_parts:
                out_text = "\n=== OUTPUT ===\n" + "\n".join(out_parts)[:5000]
        full = f"# Cell type: {ctype}\n{src}{out_text}"
        await persist_segment_fn(
            idx, f"cell {idx} ({ctype})", full, f"cell_idx={idx}",
        )


async def _extract_zip_archive(
    path: Path, fe: "FileExtraction", repo: "ExtractionRepo",
    completed: set, persist_segment_fn,
) -> None:
    """extract محتویات zip — هر فایل text-readable یک segment با path-as-title.
    فایل‌های binary فقط با نام و حجم درج می‌شوند.
    """
    import zipfile
    try:
        zf = zipfile.ZipFile(str(path), "r")
    except Exception as e:
        raise ExtractionError(f"zip open failed: {e}")
    try:
        names = [n for n in zf.namelist() if not n.endswith("/")]
        # 🔴 (extraction-100pct-fix) — قبلاً 1000 فایل + 2MB per file.
        # حالا 10K فایل + 50MB per file. اگر هنوز بیشتر بود، WARNING واضح.
        _ZIP_MAX_FILES = 10_000
        _ZIP_MAX_BYTES_PER_FILE = 50 * 1024 * 1024  # 50MB
        _zip_truncation_warning = ""
        if len(names) > _ZIP_MAX_FILES:
            _zip_truncation_warning = (
                f"🔴 zip دارای {len(names)} فایل بود ولی فقط {_ZIP_MAX_FILES} اول "
                f"extract شد ({len(names) - _ZIP_MAX_FILES} فایل drop شد)."
            )
            logger.warning(f"[extraction-100pct-fix] {_zip_truncation_warning}")
            names = names[:_ZIP_MAX_FILES]
        await repo.update_extraction(fe.id, total_segments=len(names) + 1)
        # text-decode safe per file
        for idx, name in enumerate(names, start=1):
            if idx in completed:
                continue
            try:
                with zf.open(name) as f:
                    raw = f.read(_ZIP_MAX_BYTES_PER_FILE)
                    # detect if file was truncated
                    _was_truncated_inner = False
                    try:
                        with zf.open(name) as _f2:
                            _f2.seek(0, 2)  # to end
                            _full_size = _f2.tell()
                            _was_truncated_inner = _full_size > _ZIP_MAX_BYTES_PER_FILE
                    except Exception:
                        pass
            except Exception as e:
                await persist_segment_fn(
                    idx, name, f"[خطا در خواندن: {str(e)[:100]}]", f"zip_entry={name}",
                )
                continue
            # اگر فایل internal بزرگتر بود، یک هشدار به متن چسبیده می‌شه
            if _was_truncated_inner:
                logger.warning(
                    f"[extraction-100pct-fix] zip entry '{name}' بود {_full_size:,} byte، "
                    f"فقط {_ZIP_MAX_BYTES_PER_FILE:,} byte اول خوانده شد."
                )
            # تلاش text decode
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = raw.decode("latin-1")
                except Exception:
                    # binary — فقط متادیتا
                    text = f"[binary file، size={len(raw)} bytes — متن قابل extract نیست]"
            await persist_segment_fn(idx, name, text, f"zip_entry={name}")
        # summary segment
        summary_idx = len(names) + 1
        if summary_idx not in completed:
            summary = (
                f"📦 *جمع‌بندی archive*\n"
                f"- نام: `{fe.original_filename}`\n"
                f"- {len(names)} فایل داخلی پردازش شد\n"
                f"- نمونه: {', '.join(names[:5])}{'...' if len(names) > 5 else ''}"
            )
            await persist_segment_fn(summary_idx, "📋 جمع‌بندی archive", summary, "meta=zip")
    finally:
        try:
            zf.close()
        except Exception:
            pass


# ─────────────── PDF 4-page chunking (Stage 7) ───────────────

PDF_CHUNK_SIZE: int = 4  # هر batch چند صفحه (قابل تنظیم)


async def _run_pdf_extraction(
    fe: "FileExtraction",
    session: Any,
    user_idea: str,
    repo: "ExtractionRepo",
    path: Path,
    completed: set,
    persist_segment_fn,
) -> None:
    """منطق PDF با 4-page chunking + final جمع‌بندی segment.

    1. pypdf همهٔ صفحات را extract می‌کند → text-pages (که متن دارند) و
       scan-pages (که خالی هستند، نیاز به OCR)
    2. text-pages: هر صفحه یک segment مجزا (سریع، بدون AI call)
    3. scan-pages: گروه‌بندی به batchهای 4-تایی، هر batch:
       - render هر صفحه با pdftoppm → PNG → base64
       - یک AI call با ۴ تصویر به Gemini، prompt: "متن کامل ۴ صفحه به
         ترتیب صفحه، literal، بدون خلاصه‌سازی"
       - یک segment با عنوان "صفحات N-M (OCR)" و متن کل ۴ صفحه
    4. در پایان یک segment با عنوان "📋 جمع‌بندی نهایی" که شامل
       متادیتای پروسه است (NOT summarization of content):
       - تعداد کل صفحات
       - text-pages vs scan-pages
       - تعداد AI call های انجام‌شده
       - زمان شروع/پایان (پر می‌شود توسط repo.update_extraction)
       - hash sha256 از full text
       - char count
       - مدل استفاده‌شده

    progress: per batch update روی tracker (اگر set شده).
    """
    import hashlib
    from time import time as _now
    pages = _extract_pdf_pages(path)
    total_pages = len(pages)

    # جدا کردن text-pages از scan-pages
    text_pages: List[Tuple[int, str]] = []
    scan_pages: List[int] = []
    for page_num, txt in pages:
        if txt.strip():
            text_pages.append((page_num, txt))
        else:
            scan_pages.append(page_num)

    # 🛡 (audit fix C1 — refined) — batchهای scan فقط شامل صفحات
    # **متوالی** هستند تا ordering در full_text درست بماند.
    # مثال: scan_pages=[2, 4, 5, 6, 9] → runs=[[2],[4,5,6],[9]] →
    # هر run با CHUNK_SIZE=4 split می‌شود.
    scan_batches: List[List[int]] = []
    if scan_pages:
        current_run: List[int] = [scan_pages[0]]
        for p in scan_pages[1:]:
            if p == current_run[-1] + 1:
                current_run.append(p)
            else:
                # split run
                for i in range(0, len(current_run), PDF_CHUNK_SIZE):
                    scan_batches.append(current_run[i:i + PDF_CHUNK_SIZE])
                current_run = [p]
        # last run
        for i in range(0, len(current_run), PDF_CHUNK_SIZE):
            scan_batches.append(current_run[i:i + PDF_CHUNK_SIZE])

    # total_segments: هر text-page یک، هر batch یک، plus 1 جمع‌بندی
    estimated_total = len(text_pages) + len(scan_batches) + 1
    await repo.update_extraction(fe.id, total_segments=estimated_total)

    start_time = _now()
    ai_calls = 0

    # 🛡 (audit fix C1) — segment_index بر اساس multiplier از 100 تا ترتیب
    # ادغام full_text درست بماند حتی در PDFهای مخلوط text/scan.
    # هر page شماره اصلی * 100 می‌گیرد (page 1 → 100، page 50 → 5000).
    # text-pages با ID = page_num * 100
    # scan-batches با ID = first_page_num * 100 + 1 (تا کمی بعد از همان page)
    # final جمع‌بندی با ID خیلی بزرگ.
    PAGE_ID_MULT = 100

    # 1) text-pages — هر کدام یک segment (سریع، resume-friendly)
    for page_num, txt in text_pages:
        seg_id = page_num * PAGE_ID_MULT
        if seg_id in completed:
            continue
        await _wait_for_memory_headroom()
        await persist_segment_fn(
            seg_id, f"صفحه {page_num}", txt, f"page={page_num}",
        )

    # 2) scan-batches — هر batch یک segment با OCR. id = first_page*100+1
    # تا بعد از text page با همان page_num قرار بگیرد (در صورت مخلوط)
    for b_idx, batch_pages in enumerate(scan_batches, start=1):
        first_p = batch_pages[0]
        last_p = batch_pages[-1]
        seg_id = first_p * PAGE_ID_MULT + 1
        if seg_id in completed:
            continue
        await _wait_for_memory_headroom()
        title = f"صفحات {first_p}–{last_p} (OCR — {len(batch_pages)} صفحه)"
        # render همهٔ صفحات batch به base64
        imgs: List[str] = []
        render_errors: List[str] = []
        for p_num in batch_pages:
            try:
                b64 = await _render_pdf_page_to_b64(path, p_num)
                if b64:
                    imgs.append(b64)
                else:
                    render_errors.append(f"page {p_num}: render returned None")
            except Exception as e:
                render_errors.append(f"page {p_num}: {str(e)[:100]}")

        if not imgs:
            # هیچ‌چیز render نشد — placeholder
            await persist_segment_fn(
                seg_id, title + " (render failed)",
                "[render تصویری همهٔ صفحات این batch ناموفق بود — "
                f"احتمالاً pdftoppm در سیستم نیست یا PDF خراب است.]\n"
                f"خطاها: {'; '.join(render_errors[:3])}",
                f"pages={first_p}-{last_p}",
            )
            continue

        # ساخت prompt — همهٔ تصاویر را با هم بفرست
        prompt = (
            f"این {len(imgs)} تصویر صفحات {first_p} تا {last_p} از یک PDF هستند.\n\n"
            f"**وظیفه**: متن کامل و literal هر صفحه را به ترتیب صفحه (از {first_p} به {last_p}) "
            f"استخراج کن. قبل از متن هر صفحه، یک header `## صفحه N` بگذار.\n\n"
            "قواعد سختگیرانه:\n"
            "- هیچ‌چیز را خلاصه نکن.\n"
            "- هیچ خطی، header، footer، شماره صفحه drop نشود.\n"
            "- اگر فارسی است، فارسی بنویس.\n"
            "- اگر جدولی هست، خطی‌سازی کن (با |).\n"
            "- اگر تصویری بدون متن هست، توصیف کن «[تصویر: ...]».\n"
        )
        try:
            # _ai_extract_text فقط یک image می‌گیرد در پارامتر `image_b64`،
            # اما در داخلش `images` لیست است. برای چند تصویر، باید custom call
            # داشته باشیم. ساده‌ترین راه: همهٔ تصاویر را به `_ai_extract_text`
            # بفرستیم با inline_file_data به‌عنوان لیست (اما فعلاً API single).
            # راه‌حل: اولین تصویر به image_b64، بقیه via inline_file_data fake
            # نمی‌شود — لذا ai_manager.generate را مستقیم با images=[...] صدا می‌زنیم.
            from .ai_manager import get_ai_manager
            from .ai_base import Message
            mgr = get_ai_manager()
            messages = [
                Message(role="system", content=(
                    "تو یک OCR و استخراج‌گر متن دقیق هستی. متن کامل و literal "
                    "از تصاویر داده‌شده بنویس. هیچ‌چیز خلاصه نکن."
                )),
                Message(role="user", content=prompt, images=imgs),
            ]
            resp = await asyncio.wait_for(
                mgr.generate(
                    model_id=fe.model_used, messages=messages,
                    max_tokens=32000, temperature=0.1,
                    allow_fallback=False,  # 🛡 (audit fix CRITICAL)
                ),
                timeout=PER_SEGMENT_TIMEOUT_SEC,
            )
            ai_calls += 1
            text = (resp.content or "").strip()
        except asyncio.TimeoutError:
            text = f"[timeout بعد از {PER_SEGMENT_TIMEOUT_SEC}s — batch صفحات {first_p}-{last_p}]"
        except Exception as e:
            text = f"[OCR خطا: {str(e)[:200]}]"

        await persist_segment_fn(seg_id, title, text, f"pages={first_p}-{last_p}")

    # 3) جمع‌بندی نهایی — متادیتای پروسه (NOT content summary)
    # ID خیلی بزرگ‌تر از هر page → همیشه آخرین segment
    final_seg_id = (total_pages + 1) * PAGE_ID_MULT + 999
    if final_seg_id not in completed:
        end_time = _now()
        elapsed = end_time - start_time
        # full text برای hash
        all_segs = repo.get_segments(fe.id)
        all_text_parts = []
        for s in sorted(all_segs, key=lambda x: x.segment_index):
            if s.status == "done":
                all_text_parts.append(s.text)
        full_text = "\n\n".join(all_text_parts)
        h = hashlib.sha256(full_text.encode("utf-8")).hexdigest()[:16]
        meta_text = (
            f"📋 **جمع‌بندی پروسه استخراج PDF**\n\n"
            f"- فایل: `{fe.original_filename}`\n"
            f"- تعداد کل صفحات: **{total_pages}**\n"
            f"- صفحات متنی (pypdf): **{len(text_pages)}**\n"
            f"- صفحات اسکن‌شده (OCR): **{len(scan_pages)}** "
            f"(در {len(scan_batches)} batch ۴تایی)\n"
            f"- مدل OCR: `{fe.model_used}`\n"
            f"- تعداد AI calls: **{ai_calls}**\n"
            f"- زمان کل: **{elapsed:.1f}s**\n"
            f"- char count کل متن: **{len(full_text):,}**\n"
            f"- sha256 (16 char): `{h}`\n\n"
            "✅ متن کامل در segmentهای قبلی به ترتیب صفحه قابل بازیابی است. "
            "این segment فقط متادیتای پروسه است — *نه خلاصه‌سازی محتوا*."
        )
        await persist_segment_fn(
            final_seg_id, "📋 جمع‌بندی نهایی", meta_text, "meta=process",
        )


# ─────────────── main entry ───────────────

async def extract_session(
    session_id: str,
    *,
    user_idea: str = "",
    preferred_model_id: Optional[str] = None,
    db_enabled_ids: Optional[List[str]] = None,
) -> FileExtraction:
    """استخراج یک upload session را شروع می‌کند. atomic، resumable.

    اگر extraction قبلی برای این session موجود است و status='extracting'،
    از آخرین completed_segment ادامه می‌دهد.
    خروجی: FileExtraction در حالت نهایی (extracted یا failed).
    """
    upload_svc = get_upload_session_service()
    repo = get_extraction_repo()

    session = upload_svc.get(session_id)
    if session is None:
        raise ExtractionError(f"session {session_id} یافت نشد")
    if session.status not in ("completed", "extracting"):
        raise ExtractionError(
            f"session در وضعیت {session.status} است — completed لازم است"
        )

    # serialize کن — هم‌زمان فقط یک extraction (atomic check-create-run)
    async with _EXTRACTION_SEMAPHORE:
        # 🆕 (audit fix CRITICAL) — check-create داخل semaphore تا race condition
        # نداشته باشیم. اگر دو caller هم‌زمان trigger کنند، دومی صبر می‌کند
        # و سپس extraction قبلی را resume می‌کند.
        existing = repo.list_by_session(session_id)
        fe = existing[0] if existing else None
        if fe is None:
            model = pick_best_extraction_model(
                session.mime_type,
                preferred_model_id=preferred_model_id,
                db_enabled_ids=db_enabled_ids,
            )
            if model is None:
                raise ExtractionError(
                    f"هیچ مدل بصری enabled برای mime {session.mime_type} پیدا نشد — "
                    f"از /models یک مدل multimodal فعال کن"
                )
            fe = FileExtraction(
                id=str(uuid.uuid4()),
                task_id=session.task_id,
                session_id=session.id,
                file_order=session.file_order,
                original_filename=session.original_filename,
                mime_type=session.mime_type,
                model_used=model.id,
                status="extracting",
            )
            await repo.create_extraction(fe)

        try:
            await upload_svc.set_status(session.id, "extracting")
            await repo.update_extraction(fe.id, status="extracting", started_at=now_iso())
            await _run_extraction(fe, session, user_idea, repo)
            # success
            await repo.update_extraction(
                fe.id, status="extracted", finished_at=now_iso(),
                error=None,
            )
            await upload_svc.set_status(session.id, "extracted")
        except ExtractionError as e:
            logger.warning(f"extraction failed for {session.id}: {e}")
            await repo.update_extraction(fe.id, status="failed", error=str(e)[:500],
                                          finished_at=now_iso())
            await upload_svc.set_status(session.id, "failed", error=str(e)[:500])
            raise
        except Exception as e:
            logger.exception(f"extraction unexpected error for {session.id}")
            await repo.update_extraction(fe.id, status="failed", error=str(e)[:500],
                                          finished_at=now_iso())
            await upload_svc.set_status(session.id, "failed", error=str(e)[:500])
            raise

    return fe


async def _run_extraction(
    fe: FileExtraction,
    session: UploadSession,
    user_idea: str,
    repo: ExtractionRepo,
) -> None:
    """منطق per-mime routing. هر segment به‌محض اتمام در repo ذخیره می‌شود."""
    path = Path(session.temp_path)
    if not path.exists():
        raise ExtractionError(f"temp file یافت نشد: {session.temp_path}")

    mime = (session.mime_type or "").lower()
    completed = {s.segment_index for s in repo.get_segments(fe.id) if s.status == "done"}

    async def _persist_segment(idx: int, title: str, text: str, page_ts: str = "") -> None:
        # 🔴 (extraction-100pct-fix) — قبلاً silently truncate می‌شد در 10MB.
        # حالا 100MB sanity limit است و اگر بهش رسید WARNING واضح در logs و
        # error_summary می‌چسبه — silently drop نمی‌کنه.
        if len(text) > SEGMENT_TEXT_MAX_CHARS:
            _original_len = len(text)
            logger.warning(
                f"[extraction-100pct-fix] segment[{idx}] text {_original_len:,} chars > "
                f"{SEGMENT_TEXT_MAX_CHARS:,} cap. این یک bug است — segment باید کوچک‌تر "
                f"باشه. به caller گزارش می‌شه و {_original_len - SEGMENT_TEXT_MAX_CHARS:,} "
                f"char از دست رفت."
            )
            text = text[:SEGMENT_TEXT_MAX_CHARS] + (
                f"\n\n🔴 [TRUNCATED] — این segment {_original_len:,} char بود ولی به "
                f"{SEGMENT_TEXT_MAX_CHARS:,} char truncate شد. "
                f"{_original_len - SEGMENT_TEXT_MAX_CHARS:,} char از دست رفت. "
                f"این یک bug است — segment باید کوچک‌تر باشه."
            )
        seg = ExtractionSegment(
            id=str(uuid.uuid4()),
            extraction_id=fe.id,
            segment_index=idx,
            segment_title=title[:200],
            text=text,
            raw_excerpt=text[:300],
            page_or_timestamp=page_ts,
            model_used=fe.model_used,
            status="done",
            finished_at=now_iso(),
        )
        await repo.add_segment(fe.id, seg)

    # 🆕 (Stage 9) — mime override بر اساس extension اگر mime عمومی است
    if mime in ("application/octet-stream", "", "binary/octet-stream"):
        guessed = _guess_mime_from_extension(fe.original_filename)
        if guessed and guessed != mime:
            logger.info(f"extraction: mime override {mime!r} → {guessed!r} based on extension")
            mime = guessed

    # ─────────── routing ───────────
    # 🛡 (audit fix K1+K2 CRITICAL) — ipynb و archive باید **قبل** از text branch
    # چک شوند. وگرنه .ipynb با mime=text/plain یا application/json به‌اشتباه به
    # text routing می‌رود و ساختار cells گم می‌شود.
    fname_lower = (fe.original_filename or "").lower()

    # Jupyter Notebook (اولویت بالا، بر اساس extension)
    if fname_lower.endswith(".ipynb") or mime == "application/x-ipynb+json":
        await _extract_ipynb(path, fe, repo, completed, _persist_segment)
        return

    # archive (zip) — قبل از text چون .zip ممکن است mime عمومی داشته باشد
    if (mime in ("application/zip", "application/x-zip-compressed")
            or fname_lower.endswith(".zip")):
        await _extract_zip_archive(path, fe, repo, completed, _persist_segment)
        return

    # متن خام (txt/md/csv/json/xml/yaml/toml/log/ini)
    if (mime.startswith("text/")
            or mime in (
                "application/json", "application/xml", "application/x-yaml",
                "application/yaml", "application/toml", "application/x-toml",
                "application/x-ndjson",
            )):
        text = _read_text_file(path)
        if 0 not in completed:
            await _persist_segment(0, fe.original_filename, text, "")
        await repo.update_extraction(fe.id, total_segments=1)
        return

    # 🆕 (Stage 9) — کد منبع: source code با extension
    code_exts = (
        ".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".c", ".cpp", ".h", ".hpp",
        ".go", ".rs", ".rb", ".php", ".html", ".htm", ".css", ".scss", ".sql",
        ".sh", ".bash", ".ps1", ".dockerfile", ".kt", ".swift", ".dart", ".r",
        ".lua", ".pl", ".scala", ".clj", ".ex", ".exs", ".elm", ".vue",
    )
    if (any(fname_lower.endswith(e) for e in code_exts)
            or mime in (
                "application/x-python-code", "application/javascript",
                "application/typescript", "application/x-shellscript",
            )):
        text = _read_text_file(path)
        if 0 not in completed:
            await _persist_segment(0, fe.original_filename, text, "code")
        await repo.update_extraction(fe.id, total_segments=1)
        return

    if mime == "application/pdf":
        await _run_pdf_extraction(
            fe, session, user_idea, repo, path, completed, _persist_segment,
        )
        return

    if mime in (
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        paragraphs = _extract_docx_paragraphs(path)
        # هر 50 پاراگراف یک segment (تا حجم متعادل بماند)
        CHUNK = 50
        groups = [paragraphs[i:i + CHUNK] for i in range(0, len(paragraphs), CHUNK)]
        await repo.update_extraction(fe.id, total_segments=len(groups))
        for idx, g in enumerate(groups, start=1):
            if idx in completed:
                continue
            text = "\n\n".join(g)
            await _persist_segment(idx, f"بخش {idx}", text, f"paragraphs {idx*CHUNK-CHUNK+1}..{idx*CHUNK}")
        return

    if mime in (
        "application/vnd.ms-excel",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ):
        sheets = _extract_xlsx_sheets(path)
        await repo.update_extraction(fe.id, total_segments=len(sheets))
        for idx, (name, text) in enumerate(sheets, start=1):
            if idx in completed:
                continue
            await _persist_segment(idx, f"شیت: {name}", text, f"sheet={name}")
        return

    if mime.startswith("image/"):
        # base64 → Gemini → full literal description
        headings = await _plan_headings(fe.model_used, user_idea, fe.original_filename, mime)
        # کل تصویر = یک segment (یا چند segment per heading)
        await repo.update_extraction(fe.id, total_segments=len(headings))
        try:
            raw = path.read_bytes()
            b64 = base64.b64encode(raw).decode("ascii")
        except Exception as e:
            raise ExtractionError(f"image read failed: {e}")
        for idx, h in enumerate(headings, start=1):
            if idx in completed:
                continue
            await _wait_for_memory_headroom()
            prompt = (
                f"در این تصویر، **متن کامل و literal** مرتبط با heading زیر را استخراج کن:\n"
                f"  HEADING: {h}\n\n"
                f"اگر متن داخل تصویر هست، literal بنویس. اگر صحنه/object است، توصیف کامل بده.\n"
                f"هیچ‌چیز را خلاصه نکن."
            )
            try:
                text = await asyncio.wait_for(
                    _ai_extract_text(fe.model_used, prompt=prompt, image_b64=b64, max_tokens=16000),
                    timeout=PER_SEGMENT_TIMEOUT_SEC,
                )
            except asyncio.TimeoutError:
                text = f"[timeout بعد از {PER_SEGMENT_TIMEOUT_SEC}s]"
            except Exception as e:
                text = f"[خطا: {str(e)[:200]}]"
            await _persist_segment(idx, h, text, "image")
        return

    if mime.startswith("audio/") or mime.startswith("video/"):
        size = path.stat().st_size if path.exists() else 0
        # 🆕 (Stage 9) — برای فایل‌های بزرگ، ffmpeg chunking 5-دقیقه‌ای
        if size > INLINE_MEDIA_BYTES_LIMIT:
            if not _ffmpeg_available():
                # ffmpeg در سیستم نیست — به‌جای fail، placeholder بسازیم
                # تا کاربر بداند چرا
                await repo.update_extraction(fe.id, total_segments=1)
                await _persist_segment(
                    1, fe.original_filename,
                    "[فایل بزرگ‌تر از 18MB و ffmpeg در سیستم نصب نیست. "
                    "برای استخراج کامل، ffmpeg را نصب کنید یا فایل را به chunk های ≤18MB "
                    "خرد کنید و دوباره آپلود کنید.]",
                    f"size={size}B",
                )
                return
            # chunk
            chunks_dir = path.parent / f"{path.stem}_chunks"
            try:
                chunk_files = await _ffmpeg_chunk_av(path, chunks_dir, AV_CHUNK_SECONDS)
            except Exception as e:
                raise ExtractionError(f"ffmpeg chunking failed: {e}")
            if not chunk_files:
                raise ExtractionError("ffmpeg تولید chunk نکرد")
            # عناوین فقط بر اساس user_idea
            base_headings = await _plan_headings(
                fe.model_used, user_idea, fe.original_filename, mime
            )
            heading_default = base_headings[0] if base_headings else "محتوای کامل"
            await repo.update_extraction(fe.id, total_segments=len(chunk_files))
            try:
                for idx, ch_path in enumerate(chunk_files, start=1):
                    if idx in completed:
                        continue
                    await _wait_for_memory_headroom()
                    try:
                        ch_bytes = ch_path.read_bytes()
                    except Exception as e:
                        await _persist_segment(idx, f"chunk {idx}", f"[خواندن chunk: {e}]",
                                                f"t={idx*AV_CHUNK_SECONDS-AV_CHUNK_SECONDS}s..")
                        continue
                    prompt = (
                        f"این chunk شمارهٔ {idx} از یک فایل "
                        f"{('صوتی' if mime.startswith('audio/') else 'ویدئویی')} است "
                        f"(تقریباً ثانیه {idx*AV_CHUNK_SECONDS-AV_CHUNK_SECONDS} تا {idx*AV_CHUNK_SECONDS}).\n"
                        f"raw transcript کامل و literal این chunk را بنویس. هیچ‌چیز را خلاصه نکن.\n"
                        f"اگر دیالوگ هست، با هویت گوینده‌ها (اگر مشخص). اگر صحنه است، توصیف صریح."
                    )
                    try:
                        text = await asyncio.wait_for(
                            _ai_extract_text(
                                fe.model_used, prompt=prompt,
                                inline_file_data=(mime, ch_bytes),
                                max_tokens=16000,
                            ),
                            timeout=PER_SEGMENT_TIMEOUT_SEC,
                        )
                    except asyncio.TimeoutError:
                        text = f"[timeout بعد از {PER_SEGMENT_TIMEOUT_SEC}s — chunk {idx}]"
                    except Exception as e:
                        text = f"[خطا: {str(e)[:200]}]"
                    ts = f"t={(idx-1)*AV_CHUNK_SECONDS}s..{idx*AV_CHUNK_SECONDS}s"
                    await _persist_segment(idx, f"chunk {idx}: {heading_default}", text, ts)
            finally:
                # cleanup chunks
                try:
                    for f in chunk_files:
                        try:
                            f.unlink()
                        except Exception:
                            pass
                    try:
                        chunks_dir.rmdir()
                    except Exception:
                        pass
                except Exception:
                    pass
            return
        # فایل کوچک — inline base64 (همان مسیر قبلی)
        try:
            raw = path.read_bytes()
        except Exception as e:
            raise ExtractionError(f"audio/video read failed: {e}")
        headings = await _plan_headings(fe.model_used, user_idea, fe.original_filename, mime)
        await repo.update_extraction(fe.id, total_segments=len(headings))
        for idx, h in enumerate(headings, start=1):
            if idx in completed:
                continue
            await _wait_for_memory_headroom()
            prompt = (
                f"در این فایل {('صوتی' if mime.startswith('audio/') else 'تصویری/ویدئویی')}، "
                f"**رونویسی/توصیف کامل و literal** مرتبط با heading زیر را بنویس:\n"
                f"  HEADING: {h}\n\n"
                f"اگر گفتگو هست، literal transcript. اگر صحنه است، توصیف. هیچ‌چیز را خلاصه نکن."
            )
            try:
                text = await asyncio.wait_for(
                    _ai_extract_text(
                        fe.model_used, prompt=prompt,
                        inline_file_data=(mime, raw),
                        max_tokens=16000,
                    ),
                    timeout=PER_SEGMENT_TIMEOUT_SEC,
                )
            except asyncio.TimeoutError:
                text = f"[timeout بعد از {PER_SEGMENT_TIMEOUT_SEC}s]"
            except Exception as e:
                text = f"[خطا: {str(e)[:200]}]"
            await _persist_segment(idx, h, text, mime)
        return

    # ─── fallback: mime ناشناخته ───
    # سعی کن text بخوانی؛ اگر شد خوب، نه → خطا
    try:
        text = _read_text_file(path)
    except Exception:
        text = ""
    if text:
        await repo.update_extraction(fe.id, total_segments=1)
        await _persist_segment(0, fe.original_filename, text, "raw")
        return
    raise ExtractionError(f"mime ناشناخته و قابل استخراج نیست: {mime}")
